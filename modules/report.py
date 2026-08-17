"""
Geração do relatório de auditoria em formato .docx.
Inclui gráficos de conformidade gerados com matplotlib.
"""
from __future__ import annotations

import io
import math
from datetime import date, datetime

import matplotlib
matplotlib.use("Agg")  # backend sem interface gráfica
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from modules.excel_loader import AuditData, DetalhamentoData, DistribuicaoData
from modules.state import COL_CONFORMIDADE, COL_MOTIVO, COL_ACAO

# ---------------------------------------------------------------------------
# Paleta de cores
# ---------------------------------------------------------------------------
COR_CONFORME = "#2ecc71"
COR_NAO_CONFORME = "#e74c3c"
COR_NAO_AUDITADA = "#bdc3c7"
COR_TITULO = RGBColor(0x1A, 0x3A, 0x6A)
COR_HEADER_HEX = "1A3A6A"


# ---------------------------------------------------------------------------
# Helpers de formatação Word
# ---------------------------------------------------------------------------

def _fmt_date(dt: datetime | date | None) -> str:
    if dt is None:
        return "N/D"
    return dt.strftime("%d/%m/%Y") if isinstance(dt, (datetime, date)) else str(dt)


def _titulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COR_TITULO


def _subtitulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _heading(doc: Document, texto: str, level: int = 1) -> None:
    p = doc.add_heading(texto, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = COR_TITULO


def _para(doc: Document, texto: str, bold: bool = False) -> None:
    p = doc.add_paragraph(texto)
    if bold and p.runs:
        p.runs[0].bold = True


def _set_cell_bg(cell, hex_color: str) -> None:
    ns = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    shading = parse_xml(f'<w:shd {ns} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _tabela_2col(
    doc: Document,
    linhas: list[tuple[str, str]],
    larguras: tuple[float, float] = (3.5, 3.5),
) -> None:
    table = doc.add_table(rows=len(linhas), cols=2)
    table.style = "Table Grid"
    for i, (label, valor) in enumerate(linhas):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(valor)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Inches(larguras[0])
        row.cells[1].width = Inches(larguras[1])


def _tabela_conformidade_header(table) -> None:
    hdr = table.rows[0]
    for cell in hdr.cells:
        _set_cell_bg(cell, COR_HEADER_HEX)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Gráficos matplotlib
# ---------------------------------------------------------------------------

def _grafico_barras_setores(df: pd.DataFrame, col_setor: str, total: int) -> io.BytesIO | None:
    """
    Gera gráfico de barras horizontal com a contagem e proporção por setor de destino.
    Retorna None se o DataFrame estiver vazio ou a coluna ausente.
    """
    if df is None or df.empty or col_setor not in df.columns:
        return None

    contagem = (
        df[col_setor]
        .fillna("(não informado)")
        .value_counts()
        .sort_values(ascending=True)
    )
    if contagem.empty:
        return None

    n = len(contagem)
    fig_height = max(3.5, n * 0.45)
    fig, ax = plt.subplots(figsize=(7.5, fig_height))

    bars = ax.barh(
        contagem.index.tolist(),
        contagem.values,
        color="#3498db",
        edgecolor="white",
        linewidth=1.2,
    )

    for bar, val in zip(bars, contagem.values):
        pct = val / total * 100 if total > 0 else 0.0
        ax.text(
            bar.get_width() + max(contagem.values) * 0.01,
            bar.get_y() + bar.get_height() / 2,
            f"{val} ({pct:.1f}%)",
            va="center",
            ha="left",
            fontsize=9,
        )

    ax.set_xlabel("Quantidade de tarefas distribuídas", fontsize=9)
    ax.set_title("Distribuição de Tarefas por Setor de Destino", fontsize=11,
                 fontweight="bold", color="#1A3A6A")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_xlim(0, max(contagem.values) * 1.2)
    ax.tick_params(axis="y", labelsize=9)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _grafico_pizza(
    n_conf: int,
    n_nc: int,
    n_naud: int,
    titulo: str,
) -> io.BytesIO | None:
    """Gera gráfico de pizza de conformidade. Retorna None se não há dados auditados."""
    total_aud = n_conf + n_nc
    if total_aud == 0:
        return None

    labels, sizes, colors = [], [], []
    if n_conf > 0:
        labels.append(f"Conformes\n{n_conf} ({n_conf/total_aud*100:.1f}%)")
        sizes.append(n_conf)
        colors.append(COR_CONFORME)
    if n_nc > 0:
        labels.append(f"Não Conformes\n{n_nc} ({n_nc/total_aud*100:.1f}%)")
        sizes.append(n_nc)
        colors.append(COR_NAO_CONFORME)

    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    wedges, texts = ax.pie(
        sizes, labels=labels, colors=colors, startangle=90,
        wedgeprops={"edgecolor": "white", "linewidth": 2},
    )
    for text in texts:
        text.set_fontsize(9)
    ax.set_title(titulo, fontsize=11, fontweight="bold", pad=12, color="#1A3A6A")

    if n_naud > 0:
        nota = f"* {n_naud} tarefa(s) não auditada(s) não incluída(s)"
        fig.text(0.5, 0.01, nota, ha="center", fontsize=7.5, color="#888888")

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _grafico_barras_resumo(
    total_tarefas: int,
    total_triadas: int,
    total_nao_triadas: int,
    auditadas_triadas: int,
    auditadas_nao_triadas: int,
    conf_triadas: int,
    conf_nao_triadas: int,
) -> io.BytesIO:
    """Gráfico de barras com visão geral das estatísticas."""
    categorias = [
        "Total\nProcessadas",
        "Triadas\n(automação)",
        "Não Triadas",
        "Triadas\nAuditadas",
        "Não Triadas\nAuditadas",
    ]
    valores = [total_tarefas, total_triadas, total_nao_triadas,
               auditadas_triadas, auditadas_nao_triadas]
    cores = ["#3498db", "#2ecc71", "#e67e22", "#1abc9c", "#e74c3c"]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(categorias, valores, color=cores, edgecolor="white", linewidth=1.5)

    for bar, val in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max(valores) * 0.01,
            str(val), ha="center", va="bottom", fontsize=10, fontweight="bold"
        )

    ax.set_ylabel("Quantidade de Tarefas", fontsize=10)
    ax.set_title("Visão Geral da Auditoria", fontsize=12, fontweight="bold", color="#1A3A6A")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_ylim(0, max(valores) * 1.15)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Tabelas de detalhamento
# ---------------------------------------------------------------------------

def _tabela_nao_conformidades(
    doc: Document,
    df: pd.DataFrame,
    origem_label: str,
) -> None:
    df_nc = df[df[COL_CONFORMIDADE] == "Não Conforme"].copy()
    if df_nc.empty:
        _para(doc, f"Nenhuma não conformidade identificada nas tarefas {origem_label}.")
        return

    from modules.excel_loader import COL_TAREFA, COL_NUP
    headers = ["Tarefa", "NUP", "Motivo da Não Conformidade", "Ação Corretiva"]
    table = doc.add_table(rows=1 + len(df_nc), cols=4)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, COR_HEADER_HEX)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, (_, row) in enumerate(df_nc.iterrows(), start=1):
        r = table.rows[row_idx]
        r.cells[0].text = str(row.get(COL_TAREFA, ""))
        r.cells[1].text = str(row.get(COL_NUP, ""))
        r.cells[2].text = str(row.get(COL_MOTIVO, "") or "")
        r.cells[3].text = str(row.get(COL_ACAO, "") or "")

    for row in table.rows:
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(2.8)
        row.cells[3].width = Inches(2.8)


def _tabela_relacao_auditadas(
    doc: Document,
    df: pd.DataFrame,
    colunas_extras: list[str],
) -> None:
    """Lista completa das entradas auditadas (Conformidade != 'Não auditada')."""
    from modules.excel_loader import (
        COL_TAREFA, COL_NUP,
        COL_DIST_ID, COL_DIST_NUP, COL_DIST_SETOR_ORIGEM, COL_DIST_SETOR_DESTINO,
    )

    df_aud = df[df[COL_CONFORMIDADE] != "Não auditada"].copy()
    if df_aud.empty:
        _para(doc, "Nenhuma entrada auditada.")
        return

    # Monta colunas: base (existentes, sem duplicatas) + extras + Conformidade
    seen: set = set()
    cols_base = []
    for c in [COL_TAREFA, COL_NUP, COL_DIST_ID, COL_DIST_NUP]:
        if c in df.columns and c not in seen:
            cols_base.append(c)
            seen.add(c)
    extra_cols = [c for c in colunas_extras if c in df.columns and c not in cols_base]
    cols_show = cols_base + extra_cols + [COL_CONFORMIDADE]
    headers_map = {
        COL_TAREFA: "Tarefa",
        COL_NUP: "NUP",
        COL_DIST_ID: "Id",
        COL_DIST_NUP: "NUP",
        COL_DIST_SETOR_ORIGEM: "Setor Origem",
        COL_DIST_SETOR_DESTINO: "Setor Destino",
        COL_CONFORMIDADE: "Conformidade",
    }
    headers = [headers_map.get(c, c) for c in cols_show]

    n_cols = len(cols_show)
    table = doc.add_table(rows=1 + len(df_aud), cols=n_cols)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, COR_HEADER_HEX)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    conf_col_idx = cols_show.index(COL_CONFORMIDADE)

    for row_idx, (_, row) in enumerate(df_aud.iterrows(), start=1):
        r = table.rows[row_idx]
        for col_idx, col in enumerate(cols_show):
            r.cells[col_idx].text = str(row.get(col, "") or "")
        # Colorir conformidade
        conf = str(row.get(COL_CONFORMIDADE, ""))
        if conf == "Conforme":
            _set_cell_bg(r.cells[conf_col_idx], "d5f5e3")
        elif conf == "Não Conforme":
            _set_cell_bg(r.cells[conf_col_idx], "fadbd8")


def _section_auditoria(
    doc: Document,
    numero: str,
    titulo: str,
    df: pd.DataFrame | None,
    tipo_controle: str | None,
    tamanho_amostra: int | None,
    colunas_extras: list[str],
    origem_label: str,
) -> None:
    """Renderiza seção completa de auditoria (triadas ou não-triadas)."""
    _heading(doc, f"{numero}. {titulo}")

    if df is None or df.empty:
        _para(doc, "Nenhuma tarefa selecionada para auditoria neste ciclo.")
        doc.add_paragraph()
        return

    from modules.state import stats_df
    s = stats_df(df)
    n_naud = s["total"] - s["auditadas"]

    # Subseção: resultado quantitativo
    _heading(doc, f"{numero}.1 Resultado Quantitativo", level=2)
    linhas_stats = [
        ("Total de tarefas disponíveis", str(s["total"])),
        ("Tarefas auditadas", str(s["auditadas"])),
        ("Tarefas não auditadas (excluídas das estatísticas)", str(n_naud)),
        ("Conformes", f"{s['conformes']} ({s['pct_conf']:.1f}%)"),
        ("Não Conformes", f"{s['nao_conformes']} ({s['pct_nc']:.1f}%)"),
    ]
    if tipo_controle == "detalhado" and tamanho_amostra is not None and numero == "4":
        linhas_stats.insert(1, ("Amostra definida pela fórmula", str(tamanho_amostra)))
    _tabela_2col(doc, linhas_stats)
    doc.add_paragraph()

    # Gráfico de pizza
    grafico = _grafico_pizza(s["conformes"], s["nao_conformes"], n_naud,
                             f"Conformidade — {titulo}")
    if grafico:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(grafico, width=Inches(4.0))
    doc.add_paragraph()

    # Subseção: não conformidades
    _heading(doc, f"{numero}.2 Detalhamento das Não Conformidades", level=2)
    _tabela_nao_conformidades(doc, df, origem_label)
    doc.add_paragraph()

    # Subseção: relação de auditadas
    _heading(doc, f"{numero}.3 Relação de Tarefas Auditadas", level=2)
    _tabela_relacao_auditadas(doc, df, colunas_extras)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Geração do texto de conclusão
# ---------------------------------------------------------------------------

def _conclusao(df_tri: pd.DataFrame | None, df_nao: pd.DataFrame | None) -> str:
    from modules.state import stats_df
    s_tri = stats_df(df_tri)
    s_nao = stats_df(df_nao)

    total_aud = s_tri["auditadas"] + s_nao["auditadas"]
    total_nc = s_tri["nao_conformes"] + s_nao["nao_conformes"]

    if total_aud == 0:
        return "Nenhuma tarefa foi auditada neste ciclo."

    pct_conf = (total_aud - total_nc) / total_aud * 100

    partes = [
        f"No presente ciclo de auditoria foram examinadas {total_aud} tarefa(s), "
        f"sendo {s_tri['auditadas']} tarefa(s) triada(s) e "
        f"{s_nao['auditadas']} tarefa(s) não triada(s). "
    ]

    if total_nc == 0:
        partes.append(
            "Não foram identificadas não conformidades, demonstrando adequação "
            "das regras de negócio e dos fluxos de automação do Conecta+. "
        )
    else:
        partes.append(
            f"Foram identificadas {total_nc} não conformidade(s) "
            f"({100 - pct_conf:.1f}% do total auditado). "
            "As respectivas ações corretivas foram registradas nas seções anteriores "
            "e devem ser implementadas e verificadas no próximo ciclo. "
        )

    partes.append(
        "Recomenda-se a manutenção do controle de qualidade periódico, "
        "o registro dos resultados em NUP próprio com responsáveis e periodicidade definidos, "
        "e a revisão contínua das regras de negócio, conforme preconiza o "
        "Manual de Gerenciamento Estratégico de Contencioso (Portaria PGF/AGU n. 541/2025, seção 5)."
    )

    return " ".join(partes)


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def gerar_relatorio(
    audit_data: AuditData,
    df_triadas: pd.DataFrame | None,
    df_nao_triadas: pd.DataFrame | None,
    tipo_controle: str | None,
    tamanho_amostra: int | None,
    responsavel: str,
    data_auditoria: date,
) -> bytes:
    """
    Gera o relatório de auditoria e retorna os bytes do arquivo .docx.
    Não realiza escrita em disco.
    """
    from modules.excel_loader import COL_CONFIG, COL_STATUS
    from modules.state import stats_df

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # -----------------------------------------------------------------------
    # Cabeçalho
    # -----------------------------------------------------------------------
    _titulo(doc, "RELATÓRIO DE AUDITORIA")
    _subtitulo(doc, "Conecta+ Automação — Controle de Qualidade da Triagem")
    _subtitulo(doc, "Procuradoria-Geral Federal / Advocacia-Geral da União")
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 1. Identificação
    # -----------------------------------------------------------------------
    _heading(doc, "1. IDENTIFICAÇÃO")
    periodo = (
        f"{_fmt_date(audit_data.periodo_inicio)} a {_fmt_date(audit_data.periodo_fim)}"
        if audit_data.periodo_inicio else "N/D"
    )
    _tabela_2col(doc, [
        ("Período Auditado", periodo),
        ("Data de Emissão do Relatório", _fmt_date(data_auditoria)),
        ("Responsável pela Auditoria", responsavel or "Não informado"),
        ("Sistema Auditado", "Conecta+ Automação — Módulo de Triagem Avançada"),
        ("Arquivo(s) Analisado(s)", audit_data.nome_arquivo),
        ("Base Normativa", "Portaria PGF/AGU n. 541/2025 — Manual de Gerenciamento Estratégico"),
    ])
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 2. Estatísticas Gerais da Triagem
    # -----------------------------------------------------------------------
    _heading(doc, "2. ESTATÍSTICAS GERAIS DA TRIAGEM")
    _tabela_2col(doc, [
        ("Total de Tarefas Processadas pelo Sistema", str(audit_data.total_tarefas)),
        ("Tarefas Triadas (com configurações encontradas)", f"{audit_data.total_triadas} ({audit_data.pct_triadas:.1f}%)"),
        ("Tarefas Não Triadas", f"{audit_data.total_nao_triadas} ({audit_data.pct_nao_triadas:.1f}%)"),
    ])
    doc.add_paragraph()

    # Gráfico de barras visão geral
    s_tri = stats_df(df_triadas)
    s_nao = stats_df(df_nao_triadas)
    grafico_geral = _grafico_barras_resumo(
        audit_data.total_tarefas,
        audit_data.total_triadas,
        audit_data.total_nao_triadas,
        s_tri["auditadas"],
        s_nao["auditadas"],
        s_tri["conformes"],
        s_nao["conformes"],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(grafico_geral, width=Inches(5.5))
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 3. Metodologia
    # -----------------------------------------------------------------------
    _heading(doc, "3. METODOLOGIA DO CONTROLE DE QUALIDADE")

    tipo_label = {
        "simplificado": "Controle Simplificado",
        "detalhado": "Controle Detalhado (Amostragem Estatística)",
    }.get(tipo_controle or "", tipo_controle or "Não informado")

    linhas_met = [
        ("Tipo de Controle (Tarefas Triadas)", tipo_label),
        ("Nível de Confiança", "95%"),
        ("Margem de Erro", "5%"),
    ]

    if tipo_controle == "detalhado" and tamanho_amostra is not None:
        linhas_met += [
            ("Universo Amostral", f"{audit_data.total_triadas} tarefas triadas"),
            ("Tamanho da Amostra (fórmula)", f"{tamanho_amostra} tarefas"),
            ("Fórmula Aplicada", "n = n₀ / (1 + (n₀ - 1) / N), onde n₀ = Z² · p · (1-p) / E²"),
            ("Parâmetros", "Z = 1,96 | p = 0,50 | E = 0,05"),
            ("Seleção", "Aleatória simples sem reposição"),
        ]
    if s_nao["auditadas"] > 0:
        linhas_met.append(
            ("Tarefas Não Triadas Auditadas",
             f"{s_nao['auditadas']} de {audit_data.total_nao_triadas} disponíveis")
        )

    _tabela_2col(doc, linhas_met)
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 4. Auditoria das Tarefas Triadas
    # -----------------------------------------------------------------------
    _section_auditoria(
        doc=doc,
        numero="4",
        titulo="AUDITORIA DAS TAREFAS TRIADAS",
        df=df_triadas,
        tipo_controle=tipo_controle,
        tamanho_amostra=tamanho_amostra,
        colunas_extras=[COL_CONFIG],
        origem_label="triadas",
    )

    # -----------------------------------------------------------------------
    # 5. Auditoria das Tarefas Não Triadas
    # -----------------------------------------------------------------------
    _section_auditoria(
        doc=doc,
        numero="5",
        titulo="AUDITORIA DAS TAREFAS NÃO TRIADAS",
        df=df_nao_triadas,
        tipo_controle=None,
        tamanho_amostra=None,
        colunas_extras=[COL_STATUS],
        origem_label="não triadas",
    )

    # -----------------------------------------------------------------------
    # 6. Conclusão
    # -----------------------------------------------------------------------
    _heading(doc, "6. CONCLUSÃO")
    _para(doc, _conclusao(df_triadas, df_nao_triadas))
    doc.add_paragraph()

    # Assinatura
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Brasília, {_fmt_date(data_auditoria)}").italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("_" * 50)
    r2.bold = False

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(responsavel or "Responsável pela Auditoria")
    r3.bold = True

    # -----------------------------------------------------------------------
    # Salvar em memória
    # -----------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Relatório de Auditoria de Distribuição
# ---------------------------------------------------------------------------

def gerar_relatorio_distribuicao(
    dist_data: DistribuicaoData,
    df_distribuicao: pd.DataFrame | None,
    tipo_controle: str | None,
    tamanho_amostra: int | None,
    responsavel: str,
    data_auditoria: date,
) -> bytes:
    """
    Gera o relatório de auditoria de Distribuição SS e retorna bytes do .docx.
    """
    from modules.excel_loader import (
        COL_DIST_ID, COL_DIST_NUP, COL_DIST_SETOR_DESTINO,
        COL_DIST_SETOR_ORIGEM, COL_DIST_FONTE_DADOS, COL_DIST_USUARIO_DESTINO,
    )
    from modules.state import stats_df

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _titulo(doc, "RELATÓRIO DE AUDITORIA")
    _subtitulo(doc, "Super Sapiens — Controle de Conformidade da Distribuição de Tarefas")
    _subtitulo(doc, "Procuradoria-Geral Federal / Advocacia-Geral da União")
    doc.add_paragraph()

    # 1. Identificação
    _heading(doc, "1. IDENTIFICAÇÃO")
    periodo = (
        f"{_fmt_date(dist_data.periodo_inicio)} a {_fmt_date(dist_data.periodo_fim)}"
        if dist_data.periodo_inicio else "N/D"
    )
    _tabela_2col(doc, [
        ("Período Auditado", periodo),
        ("Data de Emissão do Relatório", _fmt_date(data_auditoria)),
        ("Responsável pela Auditoria", responsavel or "Não informado"),
        ("Sistema Auditado", "Super Sapiens — Distribuição de Tarefas Judiciais"),
        ("Arquivo Analisado", dist_data.nome_arquivo),
        ("Usuário Distribuidor", dist_data.usuario_distribuidor or "N/D"),
        ("Base Normativa", "Portaria PGF/AGU n. 541/2025 — Manual de Gerenciamento Estratégico"),
    ])
    doc.add_paragraph()

    # 2. Estatísticas Gerais
    _heading(doc, "2. ESTATÍSTICAS GERAIS DA DISTRIBUIÇÃO")
    s = stats_df(df_distribuicao)
    n_amostra_real = s["total"]
    _tabela_2col(doc, [
        ("Total de Distribuições no Relatório", str(dist_data.total_distribuicoes)),
        ("Distribuições na Amostra Auditada", str(n_amostra_real)),
        ("Distribuições Auditadas", str(s["auditadas"])),
        ("Distribuições Não Auditadas (excluídas das estatísticas)", str(s["total"] - s["auditadas"])),
        ("Conformes (Setor Destino correto)", f"{s['conformes']} ({s['pct_conf']:.1f}%)"),
        ("Não Conformes", f"{s['nao_conformes']} ({s['pct_nc']:.1f}%)"),
    ])
    doc.add_paragraph()

    # Gráfico de conformidade
    grafico = _grafico_pizza(s["conformes"], s["nao_conformes"], s["total"] - s["auditadas"],
                             "Conformidade — Setor de Destino")
    if grafico:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(grafico, width=Inches(4.0))
    doc.add_paragraph()

    # 2.1 Distribuição por Setor
    _heading(doc, "2.1 Distribuição de Tarefas por Setor", level=2)
    _para(doc, (
        "A tabela e o gráfico a seguir apresentam a proporção de tarefas distribuídas "
        "para cada Setor de Destino em relação ao total de distribuições constantes do "
        "relatório exportado do Super Sapiens."
    ))
    doc.add_paragraph()

    df_full = dist_data.df
    total_dist = dist_data.total_distribuicoes
    if not df_full.empty and COL_DIST_SETOR_DESTINO in df_full.columns:
        contagem_setores = (
            df_full[COL_DIST_SETOR_DESTINO]
            .fillna("(não informado)")
            .value_counts()
            .reset_index()
        )
        contagem_setores.columns = ["Setor de Destino", "Quantidade"]
        contagem_setores["Proporção (%)"] = contagem_setores["Quantidade"].apply(
            lambda v: f"{v / total_dist * 100:.1f}%" if total_dist > 0 else "—"
        )

        # Tabela de setores
        n_rows = len(contagem_setores) + 1
        tbl = doc.add_table(rows=n_rows, cols=3)
        tbl.style = "Table Grid"
        headers_set = ["Setor de Destino", "Quantidade", "Proporção (%)"]
        for i, h in enumerate(headers_set):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, COR_HEADER_HEX)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_idx, row_data in enumerate(contagem_setores.itertuples(index=False), start=1):
            r = tbl.rows[row_idx]
            r.cells[0].text = str(row_data[0])
            r.cells[1].text = str(row_data[1])
            r.cells[2].text = str(row_data[2])
        for row in tbl.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(1.2)
            row.cells[2].width = Inches(1.5)
        doc.add_paragraph()

        # Gráfico de barras por setor
        grafico_setores = _grafico_barras_setores(df_full, COL_DIST_SETOR_DESTINO, total_dist)
        if grafico_setores:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(grafico_setores, width=Inches(6.0))
    else:
        _para(doc, "Dados de setor de destino não disponíveis.")
    doc.add_paragraph()

    # 3. Metodologia
    _heading(doc, "3. METODOLOGIA DO CONTROLE DE QUALIDADE")
    tipo_label = {
        "simplificado": "Controle Simplificado",
        "detalhado": "Controle Detalhado (Amostragem Estatística)",
    }.get(tipo_controle or "", tipo_controle or "Não informado")
    linhas_met = [
        ("Objeto de Auditoria", "Conformidade do Setor de Destino da Distribuição"),
        ("Tipo de Controle", tipo_label),
        ("Nível de Confiança", "95%"),
        ("Margem de Erro", "5%"),
    ]
    if tipo_controle == "detalhado" and tamanho_amostra:
        linhas_met += [
            ("Universo Amostral", f"{dist_data.total_distribuicoes} distribuições"),
            ("Tamanho da Amostra (fórmula)", f"{tamanho_amostra} distribuições"),
            ("Fórmula Aplicada", "n = n₀ / (1 + (n₀ - 1) / N), onde n₀ = Z² · p · (1-p) / E²"),
            ("Parâmetros", "Z = 1,96 | p = 0,50 | E = 0,05"),
            ("Seleção", "Aleatória simples sem reposição"),
        ]
    _tabela_2col(doc, linhas_met)
    doc.add_paragraph()

    _para(doc, (
        "Todas as tarefas incluídas na amostra auditada foram conferidas manualmente "
        "mediante consulta direta ao sistema Super Sapiens. O procedimento consistiu em "
        "verificar, para cada tarefa distribuída, se ela foi efetivamente encerrada no "
        "próprio setor para o qual foi distribuída ou se foi redistribuída para outro setor, "
        "o que indica erro na distribuição inicial. Consideram-se desconformidades tanto as "
        "redistribuições decorrentes de erro no setor de destino quanto os etiquetamentos "
        "equivocados ou duplicados identificados na tarefa."
    ))
    doc.add_paragraph()

    # 4. Não Conformidades
    _heading(doc, "4. DETALHAMENTO DAS NÃO CONFORMIDADES")
    _para(doc, (
        "São consideradas não conformidades: (a) redistribuições — tarefas distribuídas "
        "a um setor que não era o adequado e posteriormente redistribuídas a outro setor; "
        "e (b) etiquetamentos equivocados ou duplicados — tarefas com etiqueta incorreta "
        "ou com mais de uma etiqueta da mesma categoria atribuída."
    ))
    doc.add_paragraph()
    if df_distribuicao is not None:
        df_nc = df_distribuicao[df_distribuicao[COL_CONFORMIDADE] == "Não Conforme"].copy()
        if df_nc.empty:
            _para(doc, "Nenhuma não conformidade identificada nas distribuições auditadas.")
        else:
            cols_nc = [c for c in [COL_DIST_ID, COL_DIST_NUP, COL_DIST_SETOR_DESTINO,
                                   COL_MOTIVO, COL_ACAO] if c in df_nc.columns]
            headers_nc = {
                COL_DIST_ID: "Id",
                COL_DIST_NUP: "NUP",
                COL_DIST_SETOR_DESTINO: "Setor Destino",
                COL_MOTIVO: "Motivo da Não Conformidade",
                COL_ACAO: "Ação Corretiva",
            }
            table = doc.add_table(rows=1 + len(df_nc), cols=len(cols_nc))
            table.style = "Table Grid"
            for i, c in enumerate(cols_nc):
                cell = table.rows[0].cells[i]
                cell.text = headers_nc.get(c, c)
                _set_cell_bg(cell, COR_HEADER_HEX)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for row_idx, (_, row) in enumerate(df_nc.iterrows(), start=1):
                r = table.rows[row_idx]
                for col_idx, c in enumerate(cols_nc):
                    r.cells[col_idx].text = str(row.get(c, "") or "")
    doc.add_paragraph()

    # 5. Relação de auditadas
    _heading(doc, "5. RELAÇÃO DE DISTRIBUIÇÕES AUDITADAS")
    if df_distribuicao is not None:
        _tabela_relacao_auditadas(
            doc, df_distribuicao,
            colunas_extras=[COL_DIST_SETOR_ORIGEM, COL_DIST_SETOR_DESTINO],
        )
    doc.add_paragraph()

    # 6. Conclusão
    _heading(doc, "6. CONCLUSÃO")
    if s["auditadas"] == 0:
        _para(doc, "Nenhuma distribuição foi auditada neste ciclo.")
    else:
        total_nc = s["nao_conformes"]
        pct_conf = s["pct_conf"]
        if total_nc == 0:
            texto_conclusao = (
                f"No presente ciclo de auditoria foram examinadas {s['auditadas']} distribuição(ões), "
                f"todas com o Setor de Destino em conformidade. "
                "Não foram identificadas não conformidades."
            )
        else:
            texto_conclusao = (
                f"No presente ciclo de auditoria foram examinadas {s['auditadas']} distribuição(ões). "
                f"Foram identificadas {total_nc} não conformidade(s) ({100 - pct_conf:.1f}% do total auditado), "
                "compreendendo redistribuições decorrentes de erro no setor de destino e/ou "
                "etiquetamentos equivocados ou duplicados. "
                "As ações corretivas foram registradas na seção 4 e devem ser "
                "implementadas e verificadas no próximo ciclo de auditoria."
            )
        texto_conclusao += (
            " Recomenda-se a manutenção do controle periódico da conformidade das distribuições, "
            "o registro dos resultados em NUP próprio e a revisão contínua dos critérios de "
            "encaminhamento, conforme o Manual de Gerenciamento Estratégico de Contencioso "
            "(Portaria PGF/AGU n. 541/2025, seção 5)."
        )
        _para(doc, texto_conclusao)
    doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Brasília, {_fmt_date(data_auditoria)}").italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("_" * 50)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(responsavel or "Responsável pela Auditoria").bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Relatório de Auditoria de Detalhamento Individual
# ---------------------------------------------------------------------------

def gerar_relatorio_detalhamento(
    det_data: DetalhamentoData,
    df_detalhamento: pd.DataFrame | None,
    tipo_controle: str | None,
    tamanho_amostra: int | None,
    responsavel: str,
    data_auditoria: date,
) -> bytes:
    """
    Gera o relatório de auditoria de atividades por NUP e retorna bytes do .docx.
    """
    from modules.excel_loader import (
        COL_DET_ATIVIDADES, COL_DET_NUP, COL_DET_RESPONSAVEL,
    )
    from modules.state import stats_df

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _titulo(doc, "RELATÓRIO DE AUDITORIA")
    _subtitulo(doc, "Detalhamento Individual — Conformidade das Atividades Lançadas")
    _subtitulo(doc, "Procuradoria-Geral Federal / Advocacia-Geral da União")
    doc.add_paragraph()

    # 1. Identificação
    _heading(doc, "1. IDENTIFICAÇÃO")
    _tabela_2col(doc, [
        ("Período Auditado", det_data.meses or "N/D"),
        ("Data de Emissão do Relatório", _fmt_date(data_auditoria)),
        ("Responsável pela Auditoria", responsavel or "Não informado"),
        ("Sistema Auditado", "Power BI — Detalhamento Individual PGF"),
        ("Arquivo Analisado", det_data.nome_arquivo),
        ("Usuário Auditado", det_data.usuario or "N/D"),
        ("Unidade", det_data.unidade or "N/D"),
        ("Região", det_data.regiao or "N/D"),
        ("Base Normativa", "Portaria PGF/AGU n. 541/2025 — Manual de Gerenciamento Estratégico"),
    ])
    doc.add_paragraph()

    # 2. Metodologia
    _heading(doc, "2. METODOLOGIA")
    _para(doc, (
        "A população auditada é composta pelos NUPs em que o usuário registrou "
        f"atividades no período: {det_data.total_nups} processos, totalizando "
        f"{det_data.total_atividades} atividades lançadas. Como o relatório de origem "
        "não identifica individualmente cada atividade, a unidade de análise adotada "
        "foi o NUP: em cada processo selecionado verificou-se, no Super Sapiens, se as "
        "atividades registradas correspondem ao efetivamente praticado."
    ))
    doc.add_paragraph()

    detalhado = tipo_controle == "detalhado"
    tipo_label = {
        "simplificado": "Controle Simplificado",
        "detalhado": "Controle Detalhado (Amostragem Estatística)",
    }.get(tipo_controle or "", tipo_controle or "Não informado")

    linhas_metodo = [
        ("Objeto da Auditoria", "Conformidade das atividades lançadas por NUP"),
        ("Unidade Amostral", "NUP (processo)"),
        ("Tipo de Controle", tipo_label),
        ("Total de NUPs no Relatório", str(det_data.total_nups)),
        ("Total de Atividades Lançadas", str(det_data.total_atividades)),
    ]
    if detalhado and tamanho_amostra:
        linhas_metodo += [
            ("Universo Amostral (N)", str(det_data.total_nups)),
            ("Tamanho da Amostra (n)", str(tamanho_amostra)),
            ("Nível de Confiança", "95%"),
            ("Margem de Erro", "±5%"),
            ("Fórmula Aplicada", "n = n₀ / (1 + (n₀ - 1) / N), onde n₀ = Z² · p · (1-p) / E²"),
            ("Parâmetros", "Z = 1,96 | p = 0,50 | E = 0,05"),
            ("Seleção", "Aleatória simples sem reposição"),
        ]
    _tabela_2col(doc, linhas_metodo)
    doc.add_paragraph()

    if detalhado and tamanho_amostra:
        _para(doc, (
            f"Aplicou-se controle detalhado, com amostragem aleatória simples de "
            f"{tamanho_amostra} NUPs, para nível de confiança de 95% e margem de erro "
            "de ±5%, conforme o Anexo III do Manual."
        ))
    elif detalhado:
        _para(doc, (
            "Aplicou-se controle detalhado, com amostragem aleatória simples para "
            "nível de confiança de 95% e margem de erro de ±5%, conforme o Anexo III "
            "do Manual."
        ))
    else:
        _para(doc, (
            "Aplicou-se controle simplificado, com verificação da totalidade dos NUPs "
            "da população ou de seleção manual do auditor."
        ))
    doc.add_paragraph()

    # 3. Resultados
    _heading(doc, "3. RESULTADOS")
    s = stats_df(df_detalhamento)
    _tabela_2col(doc, [
        ("NUPs na População", str(det_data.total_nups)),
        ("NUPs na Amostra Auditada", str(s["total"])),
        ("NUPs Examinados", str(s["auditadas"])),
        ("NUPs Não Auditados (excluídos das estatísticas)", str(s["total"] - s["auditadas"])),
        ("Conformes", f"{s['conformes']} ({s['pct_conf']:.1f}%)"),
        ("Não Conformes", f"{s['nao_conformes']} ({s['pct_nc']:.1f}%)"),
    ])
    doc.add_paragraph()

    grafico = _grafico_pizza(s["conformes"], s["nao_conformes"], s["total"] - s["auditadas"],
                             "Conformidade — Atividades por NUP")
    if grafico:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(grafico, width=Inches(4.0))
    doc.add_paragraph()

    # 4. Não conformidades
    _heading(doc, "4. NÃO CONFORMIDADES IDENTIFICADAS")
    if df_detalhamento is None or s["nao_conformes"] == 0:
        _para(doc, "Não foram identificadas não conformidades na amostra examinada.")
    else:
        df_nc = df_detalhamento[df_detalhamento[COL_CONFORMIDADE] == "Não Conforme"]
        headers_nc = ["NUP", "Responsável", "Atividades", "Motivo", "Ação Corretiva"]
        tbl = doc.add_table(rows=1 + len(df_nc), cols=len(headers_nc))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers_nc):
            tbl.rows[0].cells[i].text = h
        _tabela_conformidade_header(tbl)
        for row_idx, (_, linha) in enumerate(df_nc.iterrows(), start=1):
            r = tbl.rows[row_idx]
            r.cells[0].text = str(linha.get(COL_DET_NUP, ""))
            r.cells[1].text = str(linha.get(COL_DET_RESPONSAVEL, ""))
            r.cells[2].text = str(linha.get(COL_DET_ATIVIDADES, ""))
            r.cells[3].text = str(linha.get(COL_MOTIVO, ""))
            r.cells[4].text = str(linha.get(COL_ACAO, ""))
        for row in tbl.rows:
            row.cells[0].width = Inches(1.7)
            row.cells[1].width = Inches(1.3)
            row.cells[2].width = Inches(0.8)
            row.cells[3].width = Inches(1.6)
            row.cells[4].width = Inches(1.6)
    doc.add_paragraph()

    # 5. Conclusão
    _heading(doc, "5. CONCLUSÃO")
    if s["auditadas"] == 0:
        _para(doc, "Nenhum NUP foi auditado neste ciclo.")
    else:
        if s["nao_conformes"] == 0:
            texto_conclusao = (
                f"No presente ciclo de auditoria foram examinados {s['auditadas']} NUP(s), "
                "todos com as atividades lançadas em conformidade com o efetivamente "
                "praticado. Não foram identificadas não conformidades."
            )
        else:
            texto_conclusao = (
                f"No presente ciclo de auditoria foram examinados {s['auditadas']} NUP(s). "
                f"Foram identificadas {s['nao_conformes']} não conformidade(s) "
                f"({s['pct_nc']:.1f}% do total auditado) no lançamento das atividades. "
                "As ações corretivas foram registradas na seção 4 e devem ser "
                "implementadas e verificadas no próximo ciclo de auditoria."
            )
        texto_conclusao += (
            " Recomenda-se a manutenção do controle periódico da conformidade dos "
            "lançamentos de atividades, o registro dos resultados em NUP próprio e a "
            "orientação continuada dos responsáveis, conforme o Manual de Gerenciamento "
            "Estratégico de Contencioso (Portaria PGF/AGU n. 541/2025, seção 5)."
        )
        _para(doc, texto_conclusao)
    doc.add_paragraph()

    # Assinatura
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Brasília, {_fmt_date(data_auditoria)}").italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("_" * 50)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(responsavel or "Responsável pela Auditoria").bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
